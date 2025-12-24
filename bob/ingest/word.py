"""Word document parser."""

from __future__ import annotations

from pathlib import Path

from bob.ingest.base import DocumentSection, ParsedDocument, Parser


class WordParser(Parser):
    """Parser for Word documents (.docx)."""

    extensions = [".docx"]

    def can_parse(self, path: Path) -> bool:
        """Check if this parser can handle the given file."""
        return path.suffix.lower() in self.extensions

    def parse(self, path: Path) -> ParsedDocument:
        """Parse a Word document into paragraph-based sections."""
        from docx import Document

        doc = Document(str(path))
        sections: list[DocumentSection] = []
        full_content: list[str] = []
        current_heading: str | None = None

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            full_content.append(text)

            # Check if this is a heading
            style = para.style
            style_name = style.name if style is not None else ""
            is_heading = style_name.startswith("Heading")

            if is_heading:
                current_heading = text
                # Headings also become sections
                sections.append(
                    DocumentSection(
                        content=text,
                        locator_type="paragraph",
                        locator_value={
                            "paragraph_index": para_idx,
                            "heading": current_heading,
                            "style": style_name,
                        },
                    )
                )
            else:
                sections.append(
                    DocumentSection(
                        content=text,
                        locator_type="paragraph",
                        locator_value={
                            "paragraph_index": para_idx,
                            "parent_heading": current_heading,
                            "style": style_name,
                        },
                    )
                )

        # Extract title from core properties or first heading
        title = None
        if doc.core_properties.title:
            title = doc.core_properties.title
        else:
            for section in sections:
                if "heading" in section.locator_value:
                    title = section.locator_value.get("heading")
                    break

        return ParsedDocument(
            source_path=str(path),
            source_type="word",
            content="\n\n".join(full_content),
            sections=sections,
            title=title or path.stem,
            source_date=self.get_file_date(path),
        )
