"""Tests for parser consistency across all document types.

Validates that all parsers produce consistent locator formats.
"""

from pathlib import Path

import pytest

from bob.ingest.base import DocumentSection, ParsedDocument
from bob.ingest.registry import get_parser

# Valid locator types as defined in architecture
VALID_LOCATOR_TYPES = {"heading", "page", "paragraph", "sheet", "section", "line"}


def validate_document(doc: ParsedDocument) -> list[str]:
    """Validate a parsed document for consistency.

    Returns list of validation errors (empty if valid).
    """
    errors = []

    # Document-level validation
    if not doc.source_path:
        errors.append("Document missing source_path")
    if not doc.source_type:
        errors.append("Document missing source_type")

    # Section-level validation
    for i, section in enumerate(doc.sections):
        section_errors = validate_section(section, i)
        errors.extend(section_errors)

    return errors


def validate_section(section: DocumentSection, index: int) -> list[str]:
    """Validate a document section.

    Returns list of validation errors (empty if valid).
    """
    errors = []
    prefix = f"Section {index}"

    # Required fields
    if not section.content.strip():
        errors.append(f"{prefix}: Empty content")

    if not section.locator_type:
        errors.append(f"{prefix}: Missing locator_type")
    elif section.locator_type not in VALID_LOCATOR_TYPES:
        errors.append(
            f"{prefix}: Invalid locator_type '{section.locator_type}', "
            f"must be one of {VALID_LOCATOR_TYPES}"
        )

    if not section.locator_value:
        errors.append(f"{prefix}: Missing locator_value")

    # Type-specific validation based on documented locator formats
    # See: docs/architecture.md "Locator System" table
    if section.locator_type == "heading":
        # Markdown: {heading, start_line, end_line}
        if "heading" not in section.locator_value:
            errors.append(f"{prefix}: heading locator missing 'heading' key")
        if "start_line" not in section.locator_value:
            errors.append(f"{prefix}: heading locator missing 'start_line' key")
        if "end_line" not in section.locator_value:
            errors.append(f"{prefix}: heading locator missing 'end_line' key")

    elif section.locator_type == "page":
        # PDF: {page, total_pages}
        if "page" not in section.locator_value and "start_page" not in section.locator_value:
            errors.append(f"{prefix}: page locator missing 'page' or 'start_page' key")

    elif section.locator_type == "paragraph":
        # Word: {paragraph_index, parent_heading}
        if "paragraph_index" not in section.locator_value:
            errors.append(f"{prefix}: paragraph locator missing 'paragraph_index' key")

    elif section.locator_type == "sheet":
        # Excel: {sheet_name, row_count}
        if "sheet_name" not in section.locator_value:
            errors.append(f"{prefix}: sheet locator missing 'sheet_name' key")

    elif section.locator_type == "section":
        if "name" not in section.locator_value and "section" not in section.locator_value:
            errors.append(f"{prefix}: section locator missing 'name' or 'section' key")

    return errors


class TestMarkdownParserConsistency:
    """Test markdown parser produces consistent locators."""

    def test_heading_locators_valid(self, tmp_path: Path):
        """Markdown parser produces valid heading locators."""
        md_file = tmp_path / "test.md"
        md_file.write_text(
            """# Introduction

This is the introduction section with enough words to pass validation.

## Getting Started

Here's how to get started with the project documentation.

### Prerequisites

You need to install Python and some other tools.
"""
        )

        parser = get_parser(md_file)
        doc = parser.parse(md_file)

        errors = validate_document(doc)
        assert errors == [], f"Validation errors: {errors}"

        # Check specific locator structure
        for section in doc.sections:
            assert section.locator_type == "heading"
            assert "heading" in section.locator_value
            assert "start_line" in section.locator_value
            assert "end_line" in section.locator_value
            assert isinstance(section.locator_value["start_line"], int)
            assert isinstance(section.locator_value["end_line"], int)


class TestPDFParserConsistency:
    """Test PDF parser produces consistent locators."""

    def test_page_locators_valid(self, tmp_path: Path):
        """PDF parser produces valid page locators."""
        # Create a minimal PDF for testing
        # We'll need to check if we have pypdf available
        try:
            from pypdf import PdfWriter

            pdf_file = tmp_path / "test.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=612, height=792)  # Letter size

            # Write PDF
            with open(pdf_file, "wb") as f:
                writer.write(f)

            parser = get_parser(pdf_file)
            if parser is None:
                pytest.skip("PDF parser not available")

            doc = parser.parse(pdf_file)

            # PDF may have no sections if blank, but should still validate
            if doc.sections:
                errors = validate_document(doc)
                assert errors == [], f"Validation errors: {errors}"

        except ImportError:
            pytest.skip("pypdf not installed")


class TestWordParserConsistency:
    """Test Word parser produces consistent locators."""

    def test_paragraph_locators_valid(self, tmp_path: Path):
        """Word parser produces valid paragraph locators."""
        try:
            from docx import Document

            docx_file = tmp_path / "test.docx"
            doc = Document()
            doc.add_paragraph(
                "This is the first paragraph with enough content to be indexed properly."
            )
            doc.add_paragraph("This is the second paragraph with additional content for the test.")
            doc.save(str(docx_file))

            parser = get_parser(docx_file)
            if parser is None:
                pytest.skip("Word parser not available")

            parsed = parser.parse(docx_file)

            errors = validate_document(parsed)
            assert errors == [], f"Validation errors: {errors}"

            for section in parsed.sections:
                assert section.locator_type == "paragraph"
                assert "paragraph_index" in section.locator_value
                assert isinstance(section.locator_value["paragraph_index"], int)

        except ImportError:
            pytest.skip("python-docx not installed")


class TestExcelParserConsistency:
    """Test Excel parser produces consistent locators."""

    def test_sheet_locators_valid(self, tmp_path: Path):
        """Excel parser produces valid sheet locators."""
        try:
            from openpyxl import Workbook

            xlsx_file = tmp_path / "test.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "TestSheet"
            # Add some content
            ws["A1"] = "Header1"
            ws["B1"] = "Header2"
            ws["A2"] = "Value1 with enough words for proper content validation"
            ws["B2"] = "Value2 with additional content for the test validation"
            wb.save(str(xlsx_file))

            parser = get_parser(xlsx_file)
            if parser is None:
                pytest.skip("Excel parser not available")

            doc = parser.parse(xlsx_file)

            errors = validate_document(doc)
            assert errors == [], f"Validation errors: {errors}"

            for section in doc.sections:
                assert section.locator_type == "sheet"
                assert "sheet_name" in section.locator_value

        except ImportError:
            pytest.skip("openpyxl not installed")


class TestRecipeParserConsistency:
    """Test recipe parser produces consistent locators."""

    def test_section_locators_valid(self, tmp_path: Path):
        """Recipe parser produces valid section locators."""
        recipe_file = tmp_path / "test.recipe.yaml"
        recipe_file.write_text(
            """name: Test Recipe
description: A delicious test recipe with multiple ingredients
servings: 4
ingredients:
  - name: Flour
    amount: "2 cups"
  - name: Sugar
    amount: "1 cup"
steps:
  - Mix all dry ingredients thoroughly in a bowl
  - Add wet ingredients and stir until combined
"""
        )

        parser = get_parser(recipe_file)
        assert parser is not None

        doc = parser.parse(recipe_file)

        errors = validate_document(doc)
        assert errors == [], f"Validation errors: {errors}"


class TestLocatorFormatConsistency:
    """Test that all locators follow the documented format."""

    def test_all_locators_have_required_fields(self, tmp_path: Path):
        """All parsed locators have required fields for their type."""
        # Create test files for each type
        md_file = tmp_path / "test.md"
        md_file.write_text("# Test\n\nContent with enough words for validation.")

        parser = get_parser(md_file)
        doc = parser.parse(md_file)

        for section in doc.sections:
            locator = section.locator_value

            # All locators should be dictionaries
            assert isinstance(locator, dict), f"Locator should be dict, got {type(locator)}"

            # All locators should have at least one key
            assert len(locator) > 0, "Locator should have at least one key"

            # Values should be serializable (string, int, float, bool, None)
            for key, value in locator.items():
                assert isinstance(key, str), f"Locator key should be string, got {type(key)}"
                assert isinstance(value, (str, int, float, bool, type(None), list, dict)), (
                    f"Locator value for '{key}' not serializable: {type(value)}"
                )
